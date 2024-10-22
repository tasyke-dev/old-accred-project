import csv                         #Для работы с CSV
import docx                        #Для работы с вордом
from docx import Document
from docxtpl import DocxTemplate   #Для переноса информации в ворд
#Функции
#Функция сохраняющая данные в таблицу ворд  
#Запись в CSV (дописать строчку после основного текста)
def dowrite():
    def csv_writer(data, path):
        with open(path, "a", newline='') as csv_file:
            writer = csv.writer(csv_file, delimiter=',')
            for line in data:
                writer.writerow(line)
    if __name__ == "__main__":
        data =[a1]
        path = "output1.csv"
        csv_writer(data, path)
#Функция сохраняющая данные в таблицу ворд        
def word(z1):
    doc1 = DocxTemplate("table.docx")
    context = { 'Namedisz' : a[z1][0], 'Number' : a[z1][1],
                'Trudoem' : a[z1][2], 'h_vsego' : a[z1][4]}
    doc1.render(context)
    doc1.save("table-final.docx")

    document = Document('table-final.docx')
    table = document.add_table(rows=1, cols=4, style='Table Grid')
    row = table.rows[0]
    row.cells[0].text = "{{Namedisz}}"
    row.cells[1].text = "{{Number}}"
    row.cells[2].text = "{{Trudoem}}"
    row.cells[3].text = "{{h vsego}}"
    row.cells[0].add_table(rows = 4, cols = 1)
    document.save('table.docx')
#Чтение из CSV и вывод на экран
a=[]
def read():
    def csv_reader(file_obj):
        reader = csv.reader(file_obj)
        for row in reader:
            print(" ".join(row))
            a.append(row)

    if __name__ == "__main__":
        csv_path = "output1.csv"
        with open(csv_path, "r") as f_obj:
            csv_reader(f_obj)
#Основная программа
#Создание и заполнение ключами таблицы в ворде
document = Document()
table = document.add_table(rows=1, cols=4, style='Table Grid')
row = table.rows[0]
row.cells[0].text = "{{Namedisz}}"
row.cells[1].text = "{{Number}}"
row.cells[2].text = "{{Trudoem}}"
row.cells[3].text = "{{h_vsego}}"
row.cells[0].add_table(rows = 4, cols = 1)
document.save('table.docx')            
record1={'Namedisz': 'физ-ра',
         'Number' :'Windows',
         'Trudoem': 12,
         'Zach ed' : 15,
         'h_vsego' : 27,
         'Lekcii' : 154,
         'Prakt' : 39,
         'Lab' : 45,
         'Kyrsov/projekt' : 161,
         'CPC' : 84,
         'Vid PE' : 93}
print(record1)
#Если x=1 то записывает инфу в CSV вначале(если нет) или продолжает если есть
x=1
if x==1:
    a1 = record1.keys()
    dowrite()
    a1 = record1.values()
    dowrite()
print(read())

#Перенос инфы в ворд
z=0
#x = int(input('Сколько строк всего? '))
while z != 0:
    print('Введите номер строки которую перенесут в ворде')
    z1= int(input('Из Файла: '))
    if  1<=z1<=x: 
        word(z1)
    else:
        print('Выбран не верный вариант!')
    b=int(input('Для продолжения нажмите 1, иначе будет выход из программы: '))
    if b!=1:
        z-=1
#Возвращает значения обратно из CSV в словарь, цифры останутся целыми
a1 = {}
a1.update({a[0][0]: a[1][0], a[0][1]: a[1][1], a[0][2]: int(a[1][2]),
           a[0][3]: int(a[1][3]), a[0][4]: int(a[1][4]), a[0][5]: int(a[1][5]),
           a[0][6]: int(a[1][6]), a[0][7]: int(a[1][7]),a[0][8]: int(a[1][8]),
           a[0][9]: int(a[1][9]), a[0][10]: int(a[1][10]) })
print(a1)
print(a1.keys())
print(a1.values())
