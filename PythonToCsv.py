import csv                         #Для работы с CSV
import docx                        #Для работы с вордом
from docx import Document
from docxtpl import DocxTemplate   #Для переноса информации в ворд

#Функция сохраняющая данные в таблицу ворд
def word(z2):

    doc1 = DocxTemplate("table.docx")
    context = { 'FIO' : a[z1][1], 'Predm' : a[z1][2], 'Stepen' : a[z1][3], 'Kab' : a[z1][4]}
    doc1.render(context)
    doc1.save("table-final.docx")

    document = Document('table-final.docx')
    table = document.add_table(rows=1, cols=4, style='Table Grid')
    row = table.rows[0]
    row.cells[0].text = "{{FIO}}"
    row.cells[1].text = "{{Predm}}"
    row.cells[2].text = "{{Stepen}}"
    row.cells[3].text = "{{Kab}}"
    row.cells[0].add_table(rows = 4, cols = 1)
    document.save('table.docx')

#Запись в CSV 
g=int(input('Записать или перезаписать таблицу?(Введите 1)'))
if g == 1:
    def csv_writer(data, path):
        with open(path, "w", newline='') as csv_file:
            writer = csv.writer(csv_file, delimiter=',')
            for line in data:
                writer.writerow(line)
    if __name__ == "__main__":
        data = ["№,ФИО,Предмет,Уч.степень,Кабинет".split(","),
                "1,Пётр Перов,информатика,Доцент,А262".split(","),
                "2,Владимир Владимирович,бокс,Старший преподаватель,Б104".split(","),
                "3,Денис Ильевич,физика,Доцент,В303".split(","),
                "4,Семён Дмитриевич,математика,Профессор,Г415".split(","),
                "5,Евгений Кондратович,жизнь,Доцент,Д101".split(",")
                ]  
        path = "output.csv"
        csv_writer(data, path)

z=int(input('Для добавления введите 1: '))
while z == 1:
    a1 = str(input("Введите '№,ФИО,Предмет,Уч.степень,Кабинет': "))
    def csv_writer(data, path):
        with open(path, "a", newline='') as csv_file:
            writer = csv.writer(csv_file, delimiter=',')
            for line in data:
                writer.writerow(line)
    if __name__ == "__main__":
        data =[a1.split(",")]
        path = "output.csv"
        csv_writer(data, path)
    z=int(input('Для продолжения нажмите 1, иначе будет выход из программы'))

a = []
#Чтение из CSV и вывод на экран
def csv_reader(file_obj):
    reader = csv.reader(file_obj)
    for row in reader:
        print(" ".join(row))
        a.append(row)

if __name__ == "__main__":
    csv_path = "output.csv"
    with open(csv_path, "r") as f_obj:
        csv_reader(f_obj)
        
#Создание и заполнение ключами таблицы в ворде
    document = Document()
    table = document.add_table(rows=1, cols=4, style='Table Grid')
    row = table.rows[0]
    row.cells[0].text = "{{FIO}}"
    row.cells[1].text = "{{Predm}}"
    row.cells[2].text = "{{Stepen}}"
    row.cells[3].text = "{{Kab}}"
    row.cells[0].add_table(rows = 4, cols = 1)
    document.save('table.docx')


x= int(a[-1][0])+1
#Перенос инфы в ворд
z=1
while z != 0:
    print('Введите номер строки из экселя которую перенесут в ворде')
    z1= int(input('Из экселя: '))
    if  1<=z1<x:     
        word(z1)
    else:
        print('Выбран не верный вариант!')
    b=int(input('Для продолжения нажмите 1, иначе будет выход из программы: '))
    if b!=1:
        z-=1
