"""!!!Файл с функциями, для создания!!!"""

import docx
import pandas as pd
import numpy as np

doc = docx.Document()
df = pd.DataFrame()

#Создать заголовок
def Title():
    TextH="Тестовый заголовок"
    StyleH=7 
    doc.add_heading(TextH, StyleH)
    

#Новый абзац и добавление текста
def Paragraph():
    ParName="Тестовый абзац "
    DocName = doc.add_paragraph(ParName)
    addText="Какой-то текст "
    DocName.add_run(addText)
    

#Создание таблицы
def Tablet():
    row = 6
    col = 4
    table = doc.add_table(rows=row, cols=col)
    table.style = 'Table Grid'

#Объединение всех функций для демонстрации возможностей
def allInOne(directory):
    
    Title()
    Paragraph()
    Tablet()
    doc.save(directory)



