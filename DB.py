import openpyxl                    #Библиотека для работы с Экселем
from openpyxl import Workbook      #Для создания
from openpyxl import load_workbook #Для загрузки
import docx                        #Для работы с вордом
from docx import Document
from docxtpl import DocxTemplate   #Для переноса информации в ворд

#Функция поиска и вывода строки из экселя
def search (x1):
    if (x1 < a) and (x1 > 0):
        x1 = x1-1
        return (dict['a1'][x1],dict['a2'][x1],dict['a3'][x1],dict['a4'][x1],dict['a5'][x1])
    else:
        return ('Такого нет!')
#Функция записывающая столбец в элемент
def sbor (x, y, z):
    for cellObj in sheet[ y : z]:
      for cell in cellObj:
          x.append(cell.value)

#Функция сохраняющая данные в таблицу ворд
def word(z2):
    z2 -= 1

    doc1 = DocxTemplate("D:\\GitHub\\Accreditation\\TESST.docx")
    context = { 'FIO' : dict['a2'][z2], 'Predm' : dict['a3'][z2], 'Stepen' : dict['a4'][z2], 'Kab' : dict['a5'][z2]}
    doc1.render(context)
    doc1.save("table-final.docx")

    document = Document('table-final.docx')
    table = document.add_table(rows=1, cols=4, style='Table Grid')
    row = table.rows[0]
    row.cells[0].text = "{{FIO}}"
    row.cells[1].text = "{{Predm}}"
    row.cells[2].text = "{{Stepen}}"
    row.cells[3].text = "{{Kab}}"
    row.cells[0].add_table(rows = 4, cols = 1)
    document.save('TESST.docx')


#Загружаем файл
wb = load_workbook('D:\\GitHub\\Accreditation\\sample.xlsx')
sheet = wb.get_sheet_by_name('Sheet')
# = int(input('Сколько строк? '))
a=5
a = a + 1
b = 'E' + str(a)
#Вывести на экран
x=int(input('Вывести значения на экран?(Введите 1):'))
if x == 1:
    for cellObj in sheet['A1': b]:
        for cell in cellObj:
            print(cell.value)
        print('--- Конец строки ---')

#Сбор данных по отдельным массивам словарь
dict = {'a1':[], 'a2':[], 'a3':[], 'a4':[], 'a5':[]}
ba = 'A' + str(a)
sbor(dict['a1'], 'A2', ba)
ba = 'B' + str(a)
sbor(dict['a2'], 'B2', ba)
ba = 'C' + str(a)
sbor(dict['a3'], 'C2', ba)
ba = 'D' + str(a)
sbor(dict['a4'], 'D2', ba)
ba = 'E' + str(a)
sbor(dict['a5'], 'E2', ba)

#Вывод определённого столбца на экран 

z=1
while z != 0:
    print('Выберите какой столбик вывести на экран: ')
    print('1 – №')
    print('2 – ФИО')
    print('3 – Предмет')
    print('4 – Уч.степень')
    print('5 – Кабинет')
    x = int(input('Ваш выбор – '))
    if  1<=x<=5:
               if x==1:
                    print(dict['a1'])
                    x1 = int(input('Вывести конкретного?(цифра) '))
                    print (search(x1))
               if x==2:
                    print(dict['a2'])
                    x1 = int(input('Вывести конкретного?(цифра) '))
                    print (search(x1))
               if x==3:
                    print(dict['a3'])
                    x1 = int(input('Вывести конкретного?(цифра) '))
                    print (search(x1))
               if x==4:
                    print(dict['a4'])
                    x1 = int(input('Вывести конкретного?(цифра) '))
                    print (search(x1))
               if x==5:
                    print(dict['a5'])
                    x1 = int(input('Вывести конкретного?(цифра) '))
                    print (search(x1))
    else:
       print('Выбран не верный вариант!')
    print('Для продолжения нажмите 1, иначе будет выход из программы')
    b=int(input())
    if b!=1:
        z-=1

    document = Document()
    table = document.add_table(rows=1, cols=4, style='Table Grid')
    row = table.rows[0]
    row.cells[0].text = "{{FIO}}"
    row.cells[1].text = "{{Predm}}"
    row.cells[2].text = "{{Stepen}}"
    row.cells[3].text = "{{Kab}}"
    row.cells[0].add_table(rows = 4, cols = 1)
    document.save('D:\\GitHub\\Accreditation\\TESST.docx')

z3=1
while z3 != 0:
    print('Введите номер строки из экселя и каким он будет в ворде')
    z2= int(input('Из экселя: '))
    if  1<=z2<a:
        word(z2)
    else:
        print('Выбран не верный вариант!')
    b=int(input('Для продолжения нажмите 1, иначе будет выход из программы: '))
    if b!=1:
        z3-=1
