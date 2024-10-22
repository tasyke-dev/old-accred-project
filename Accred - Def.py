import docx
import pandas as pd
import numpy as np

doc = docx.Document('D:\\GitHub\\Accreditation\\Test.docx')
df = pd.DataFrame()

###Создать заголовок
def Title():
    TextH=input("Введите заголовок: ")
    StyleH=int(input("Стиль: "))
    doc.add_heading(TextH, StyleH)
#Title()

###Новый абзац и добавление текста
def Paragraph():
    ParName=input("Введите текст: ")
    DocName = doc.add_paragraph(ParName)
    addText=input('Добавьте текст: ')
    DocName.add_run(addText)
Paragraph()

###Создание таблицы
def Tablet():
    row = int(input('Строки: '))
    col = int(input('Столбцы: '))
    table = doc.add_table(rows=row, cols=col)
    table.style = 'Table Grid'
Tablet()

###Замена текста
def WordChange():
    for paragraph in doc.paragraphs:
        paragraph.text = paragraph.text.replace("ист", "Такое вот предложение")


WordChange()

###Замена текста в таблице
def TableChange():
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.text = paragraph.text.replace("fffffffffffffffff", "123")

TableChange()

doc.save('D:\\GitHub\\Accreditation\\Test.docx')



